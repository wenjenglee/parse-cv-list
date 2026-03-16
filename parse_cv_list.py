#!/usr/bin/env python3
"""
CV Conference List Parser
解析心血管討論會 CV list 檔案 (.docx / .pdf / .doc)
輸出：更新 Excel 總表 + JSON 供 Notion 匯入

Usage:
    python parse_cv_list.py "CV list 20260218.docx"
    python parse_cv_list.py "CV list 20260218.docx" --update-excel
    python parse_cv_list.py --batch              # 解析資料夾內所有檔案
    python parse_cv_list.py --stats              # 顯示統計摘要
"""

import argparse
import glob
import json
import os
import re
import struct
import sys
from collections import Counter
from pathlib import Path

# --- Parsers ---

def extract_date_from_filename(fname: str) -> str:
    """從檔名提取討論會日期"""
    m = re.search(r'(\d{8})', fname)
    if m:
        d = m.group(1)
        return f"{d[:4]}/{d[4:6]}/{d[6:8]}"
    m2 = re.search(r'(\d{6})', fname)
    if m2:
        d = m2.group(1)
        return f"{d[:4]}/{d[4:6]}/01"
    return "unknown"


def parse_docx(filepath: str, conf_date: str) -> list[dict]:
    """解析 .docx 檔案（Word 表格格式）"""
    from docx import Document
    doc = Document(filepath)
    # If filename date is incomplete (e.g. 202212 → 2022/12/01), try to find
    # the actual date from document paragraphs (e.g. "CV Conference 2022/12/21")
    if conf_date.endswith('/01'):
        for para in doc.paragraphs:
            m = re.search(r'(\d{4}/\d{2}/\d{2})', para.text)
            if m:
                conf_date = m.group(1)
                break
    rows_out = []
    for table in doc.tables:
        for j, row in enumerate(table.rows):
            if j == 0:
                continue
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) < 7 or not any(c for c in cells):
                continue

            name_chart = cells[0]
            parts = re.split(r'\n', name_chart)
            if len(parts) >= 2:
                p1, p2 = parts[0].strip(), parts[1].strip()
                if re.match(r'^\d+$', p1):
                    chart_no, name = p1, p2
                else:
                    name, chart_no = p1, p2
            else:
                # Name and chart no on same line: split at trailing alphanumeric ID
                m = re.match(r'^(.*?)\s*([A-Z]?\d{6,7})\s*$', name_chart.strip())
                if m:
                    name, chart_no = m.group(1).strip(), m.group(2).strip()
                else:
                    name, chart_no = name_chart, ''

            ag = cells[1].strip()
            # Handle formats: "65M", "65 M", "76 y/r\nmale", "76y/o Male"
            # Try digit+M/F pattern first (most common: "65M", "65 M")
            age_gender_m = re.match(r'(\d{2,3})\s*([MF])\b', ag, re.IGNORECASE)
            if age_gender_m:
                age = age_gender_m.group(1)
                gender = age_gender_m.group(2).upper()
            else:
                age_num = re.search(r'(\d{2,3})', ag)
                age = age_num.group(1) if age_num else ''
                gender_word = re.search(r'\b(male|female)\b', ag, re.IGNORECASE)
                if gender_word:
                    gender = 'M' if gender_word.group(1).lower() == 'male' else 'F'
                else:
                    gender = ''

            rows_out.append({
                'conference_date': conf_date,
                'name': name,
                'chart_no': chart_no,
                'age': age,
                'gender': gender,
                'reason_of_mpi': cells[2],
                'risk_factors': cells[3],
                'mpi_dates': cells[4],
                'cta_dates': cells[5],
                'cag_dates': cells[6],
                'source_file': os.path.basename(filepath),
                'data_quality': 'Complete',
            })
    return rows_out


def parse_doc_utf16(filepath: str, conf_date: str) -> list[dict]:
    """解析 .doc 檔案（透過 OLE UTF-16LE 提取）"""
    import olefile
    ole = olefile.OleFileIO(filepath)
    word_stream = ole.openstream('WordDocument').read()
    text = word_stream.decode('utf-16-le', errors='ignore')
    ole.close()

    rows_out = []
    for m in re.finditer(r'(\d{7})', text):
        chart_no = m.group(1)
        pos = m.start()
        before = text[max(0, pos - 200):pos]
        after = text[pos:pos + 800]

        age_match = re.search(r'(\d{2,3})\s*[/y]?\s*o?\s*([MF])', after)
        if not age_match:
            continue

        dates = re.findall(r'(\d{4}/\d{2}/\d{2})', after[:600])
        chinese_chars = re.findall(r'([\u4e00-\u9fff]+)', before[-60:] + after[:30])
        name = next((cn for cn in chinese_chars if 1 < len(cn) <= 4), '')

        reason = ''
        after_age = after[age_match.end():]
        reason_match = re.search(
            r'([A-Za-z][A-Za-z\s,/\(\)]+?)(?=(?:HTN|DM|DLP|Dyslipidemia|Hypertension|Age|Smoking|Obesity|gender|risk))',
            after_age
        )
        if reason_match:
            reason = reason_match.group(1).strip()

        rows_out.append({
            'conference_date': conf_date,
            'name': name,
            'chart_no': chart_no,
            'age': age_match.group(1),
            'gender': age_match.group(2),
            'reason_of_mpi': reason or '(needs review)',
            'risk_factors': '(needs review)',
            'mpi_dates': ', '.join(dates[:3]) if dates else '',
            'cta_dates': '',
            'cag_dates': '',
            'source_file': os.path.basename(filepath),
            'data_quality': 'Partial',
        })
    return rows_out


def parse_pdf(filepath: str, conf_date: str) -> list[dict]:
    """解析 PDF 檔案（heuristic text extraction）"""
    import fitz
    doc = fitz.open(filepath)
    full_text = "".join(page.get_text() for page in doc)
    lines = [l.strip() for l in full_text.split('\n') if l.strip()]

    rows_out = []
    i = 0
    while i < len(lines):
        if re.match(r'^\d{6,7}$', lines[i]):
            chart_no = lines[i]
            name = lines[i + 1] if i + 1 < len(lines) else ''
            ag_line = lines[i + 2] if i + 2 < len(lines) else ''
            age_m = re.match(r'(\d+)\s*([MF])', ag_line)
            rows_out.append({
                'conference_date': conf_date,
                'name': name.replace(' ', ''),
                'chart_no': chart_no,
                'age': age_m.group(1) if age_m else ag_line,
                'gender': age_m.group(2) if age_m else '',
                'reason_of_mpi': '(from PDF - needs review)',
                'risk_factors': '(from PDF - needs review)',
                'mpi_dates': '', 'cta_dates': '', 'cag_dates': '',
                'source_file': os.path.basename(filepath),
                'data_quality': 'Needs Review',
            })
            i += 3
        else:
            i += 1
    return rows_out


def parse_file(filepath: str) -> list[dict]:
    """自動偵測檔案格式並解析"""
    conf_date = extract_date_from_filename(os.path.basename(filepath))
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.docx':
        return parse_docx(filepath, conf_date)
    elif ext == '.doc':
        return parse_doc_utf16(filepath, conf_date)
    elif ext == '.pdf':
        return parse_pdf(filepath, conf_date)
    else:
        print(f"  [SKIP] Unsupported format: {ext}")
        return []


# --- Output ---

def build_notion_page(r: dict) -> dict:
    """將解析結果轉為 Notion page properties"""
    rf = r.get('risk_factors', '').upper()
    risk_tags = []
    if 'HTN' in rf or 'HYPERTENSION' in rf: risk_tags.append('HTN')
    if 'DM' in rf or 'DIABETES' in rf: risk_tags.append('DM')
    if 'DLP' in rf or 'DYSLIPIDEMIA' in rf or 'HYPERLIPIDEMIA' in rf or 'HLP' in rf: risk_tags.append('DLP')
    if 'SMOK' in rf: risk_tags.append('Smoking')
    if 'OBES' in rf or 'OVERWEIGHT' in rf: risk_tags.append('Obesity')
    if 'AGE' in rf: risk_tags.append('Age')
    if 'GENDER' in rf: risk_tags.append('Gender')

    conf_date = r.get('conference_date', '').replace('/', '-')
    if len(conf_date) == 7:
        conf_date += '-01'

    name = r.get('name', '').strip()
    chart = r.get('chart_no', '').strip()
    case_title = f"{chart} {name}".strip() or f"Case {conf_date}"

    props = {
        "Case": case_title,
        "Chart No": chart,
        "Reason of MPI": r.get('reason_of_mpi', '').replace('\r', ' '),
        "Risk Factors": json.dumps(risk_tags),
        "MPI Dates": r.get('mpi_dates', ''),
        "CTA Dates": r.get('cta_dates', ''),
        "CAG Dates": r.get('cag_dates', ''),
        "Data Quality": r.get('data_quality', 'Complete'),
        "Source File": r.get('source_file', ''),
    }
    try:
        props["Age"] = int(r.get('age', ''))
    except (ValueError, TypeError):
        pass
    if r.get('gender', '') in ('M', 'F'):
        props["Gender"] = r['gender']
    if conf_date and len(conf_date) >= 10:
        props["date:Conference Date:start"] = conf_date[:10]
        props["date:Conference Date:is_datetime"] = 0

    return {"properties": props}


def sanitize(val):
    """Remove illegal Excel characters (control chars except tab/newline)"""
    if not isinstance(val, str):
        return val
    import re as _re
    return _re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', val)


def update_excel(records: list[dict], excel_path: str):
    """新增記錄到 Excel 總表"""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        ws = wb['CV Conference Cases']
        start_row = ws.max_row + 1
        existing_count = ws.max_row - 1
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'CV Conference Cases'
        headers = ['No.', 'Conference Date', 'Chart No.', 'Name', 'Age', 'Gender',
                   'Reason of MPI', 'Risk Factors', 'MPI Dates', 'CTA Dates',
                   'CAG Dates', 'Source File', 'Data Quality']
        header_font = Font(name='Arial', bold=True, size=11, color='FFFFFF')
        header_fill = PatternFill('solid', fgColor='2F5496')
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        ws.freeze_panes = 'A2'
        start_row = 2
        existing_count = 0

    data_font = Font(name='Arial', size=10)
    thin_border = Border(
        left=openpyxl.styles.Side(style='thin', color='D9D9D9'),
        right=openpyxl.styles.Side(style='thin', color='D9D9D9'),
        top=openpyxl.styles.Side(style='thin', color='D9D9D9'),
        bottom=openpyxl.styles.Side(style='thin', color='D9D9D9')
    )
    alt_fill = PatternFill('solid', fgColor='F2F6FC')

    for idx, r in enumerate(records):
        row_num = start_row + idx
        values = [
            existing_count + idx + 1,
            sanitize(r.get('conference_date', '')),
            sanitize(r.get('chart_no', '')),
            sanitize(r.get('name', '')),
            r.get('age', ''),
            sanitize(r.get('gender', '')),
            sanitize(r.get('reason_of_mpi', '')),
            sanitize(r.get('risk_factors', '')),
            sanitize(r.get('mpi_dates', '')),
            sanitize(r.get('cta_dates', '')),
            sanitize(r.get('cag_dates', '')),
            sanitize(r.get('source_file', '')),
            sanitize(r.get('data_quality', '')),
        ]
        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row_num, column=col, value=val)
            cell.font = data_font
            cell.border = thin_border
            cell.alignment = Alignment(vertical='top', wrap_text=True)
            if row_num % 2 == 0:
                cell.fill = alt_fill

    ws.auto_filter.ref = f'A1:M{start_row + len(records) - 1}'
    wb.save(excel_path)
    print(f"  Excel updated: {len(records)} rows added → {excel_path}")


def show_stats(folder: str):
    """顯示資料夾統計摘要"""
    files = sorted(glob.glob(os.path.join(folder, 'CV list *')))
    files = [f for f in files if f.endswith(('.doc', '.docx', '.pdf'))]
    print(f"Total files: {len(files)}")

    all_records = []
    for f in files:
        records = parse_file(f)
        all_records.extend(records)
        print(f"  {os.path.basename(f)}: {len(records)} cases")

    dates = sorted(set(r['conference_date'] for r in all_records))
    genders = Counter(r['gender'] for r in all_records)
    quality = Counter(r.get('data_quality', '') for r in all_records)

    print(f"\n=== Summary ===")
    print(f"Total cases: {len(all_records)}")
    print(f"Conference sessions: {len(dates)}")
    print(f"Date range: {dates[0]} ~ {dates[-1]}")
    print(f"Gender: M={genders.get('M', 0)}, F={genders.get('F', 0)}")
    print(f"Quality: {dict(quality)}")


# --- Main ---

def main():
    parser = argparse.ArgumentParser(description='CV Conference List Parser')
    parser.add_argument('file', nargs='?', help='CV list file to parse')
    parser.add_argument('--update-excel', action='store_true', help='Append to Excel master file')
    parser.add_argument('--batch', action='store_true', help='Parse all files in input folder')
    parser.add_argument('--new', action='store_true', help='Parse only files not yet in Excel (auto-detect new)')
    parser.add_argument('--stats', action='store_true', help='Show statistics summary')
    parser.add_argument('--json-out', help='Output JSON file for Notion import')
    parser.add_argument('--excel-path', default=None, help='Excel file path (default: data/CV_Conference_Database.xlsx)')
    parser.add_argument('--data-dir', default=None, help='Data directory (default: data/ if exists, else script dir)')
    args = parser.parse_args()

    script_dir = os.path.dirname(os.path.abspath(__file__))
    # Resolve data directory: prefer data/ subfolder if it exists
    if args.data_dir:
        data_dir = os.path.abspath(args.data_dir)
    elif os.path.isdir(os.path.join(script_dir, 'data')):
        data_dir = os.path.join(script_dir, 'data')
    else:
        data_dir = script_dir

    # Resolve Excel path: prefer data/ location
    if args.excel_path:
        excel_path = args.excel_path
    elif os.path.exists(os.path.join(data_dir, 'CV_Conference_Database.xlsx')):
        excel_path = os.path.join(data_dir, 'CV_Conference_Database.xlsx')
    else:
        excel_path = os.path.join(script_dir, 'CV_Conference_Database.xlsx')

    if args.stats:
        show_stats(data_dir)
        return

    if args.batch:
        files = sorted(glob.glob(os.path.join(data_dir, 'CV list *')))
        files = [f for f in files if f.endswith(('.doc', '.docx', '.pdf'))]
    elif args.new:
        # Find files not yet recorded in Excel
        all_files = sorted(glob.glob(os.path.join(data_dir, 'CV list *')))
        all_files = [f for f in all_files if f.endswith(('.doc', '.docx', '.pdf'))]
        if os.path.exists(excel_path):
            import openpyxl
            wb = openpyxl.load_workbook(excel_path)
            ws = wb['CV Conference Cases']
            processed = set(ws.cell(row=r, column=12).value for r in range(2, ws.max_row + 1))
            files = [f for f in all_files if os.path.basename(f) not in processed]
            print(f"Already processed: {len(processed)} files")
            print(f"New files found: {len(files)}")
        else:
            files = all_files
            print(f"No Excel found, treating all {len(files)} files as new")
        if not files:
            print("No new files to process.")
            sys.exit(0)
        for f in files:
            print(f"  NEW: {os.path.basename(f)}")
    elif args.file:
        # Check absolute path, then data_dir, then script_dir
        if os.path.isabs(args.file):
            fpath = args.file
        elif os.path.exists(os.path.join(data_dir, args.file)):
            fpath = os.path.join(data_dir, args.file)
        else:
            fpath = os.path.join(script_dir, args.file)
        if not os.path.exists(fpath):
            print(f"File not found: {fpath}")
            sys.exit(1)
        files = [fpath]
    else:
        parser.print_help()
        sys.exit(1)

    all_records = []
    for f in files:
        print(f"Parsing: {os.path.basename(f)}")
        records = parse_file(f)
        print(f"  → {len(records)} cases extracted")
        all_records.extend(records)

    if not all_records:
        print("No records found.")
        sys.exit(0)

    # JSON output for Notion
    json_out = args.json_out or os.path.join(script_dir, 'notion_import.json')
    notion_pages = [build_notion_page(r) for r in all_records]
    with open(json_out, 'w', encoding='utf-8') as f:
        json.dump(notion_pages, f, ensure_ascii=False, indent=2)
    print(f"\nNotion JSON: {len(notion_pages)} pages → {json_out}")

    # Excel
    if args.update_excel:
        update_excel(all_records, excel_path)

    print(f"\nDone. {len(all_records)} total records processed.")


if __name__ == '__main__':
    main()
